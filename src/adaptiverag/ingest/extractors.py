from pathlib import Path
from .models import Document
import csv
import io
import re

from bs4 import BeautifulSoup
from pypdf import PdfReader
from docx import Document as DocxDocument


def extract_text(filepath: Path) -> list[Document]:
    """Extract from .txt files."""
    text = filepath.read_text(encoding="utf-8")
    return [Document(
        text=text,
        metadata={"source": str(filepath), "filetype": "txt"},
    )]


def extract_markdown(filepath: Path) -> list[Document]:
    """Extract from .md files."""
    text = filepath.read_text(encoding="utf-8")
    return [Document(
        text=text,
        metadata={"source": str(filepath), "filetype": "md"},
    )]


def extract_csv(filepath: Path) -> list[Document]:
    """Extract from .csv files. Each row becomes readable text."""
    text = filepath.read_text(encoding="utf-8")
    reader = csv.DictReader(io.StringIO(text))

    rows = []
    for i, row in enumerate(reader, start=1):
        row_text = " | ".join(f"{k}: {v}" for k, v in row.items())
        rows.append(row_text)

    return [Document(
        text="\n".join(rows),
        metadata={
            "source": str(filepath),
            "filetype": "csv",
            "row_count": len(rows),
        },
    )]


def extract_html(filepath: Path) -> list[Document]:
    """Extract visible text from .html files, stripping tags."""
    raw = filepath.read_text(encoding="utf-8")
    soup = BeautifulSoup(raw, "html.parser")

    # Remove script and style elements
    for tag in soup(["script", "style"]):
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)
    return [Document(
        text=text,
        metadata={
            "source": str(filepath),
            "filetype": "html",
            "title": soup.title.string if soup.title else None,
        },
    )]


def _strip_pdf_boilerplate(page_texts: list[str], 
                           candidate_lines: int = 5,
                           threshold: float = 0.5) -> list[str]:
    """
    Remove repeated header/footer lines from PDF pages.

    Looks at the first and last `candidate_lines` lines of each page,
    counts how often each appears across all pages, and strips any
    line found on more than `threshold` fraction of pages.
    Also strips common page-number patterns (e.g. "Page 3 of 10", "- 7 -").

    Args:
        page_texts:      Raw extracted text for each page.
        candidate_lines: How many lines to check from top/bottom of each page.
        threshold:       Fraction of pages a line must appear on to be
                         considered boilerplate (0.5 = 50%).

    Returns:
        Cleaned text for each page, same length as input.
    """
    if len(page_texts) < 3:
        # Too few pages to detect repetition reliably
        return page_texts

    # --- Step 1: collect candidate lines from top/bottom of each page ---
    line_page_count: dict[str, int] = {}

    for page_text in page_texts:
        lines = page_text.splitlines()
        top = lines[:candidate_lines]
        bottom = lines[-candidate_lines:] if len(lines) > candidate_lines else []
        candidates = top + bottom

        # Use a set so each line is counted once per page
        seen_on_this_page: set[str] = set()
        for line in candidates:
            normalized = line.strip().lower()
            if normalized and normalized not in seen_on_this_page:
                seen_on_this_page.add(normalized)
                line_page_count[normalized] = line_page_count.get(normalized, 0) + 1

    # --- Step 2: identify boilerplate lines ---
    num_pages = len(page_texts)
    min_appearances = int(num_pages * threshold)
    boilerplate: set[str] = {
        line for line, count in line_page_count.items()
        if count >= min_appearances
    }

    # --- Step 3: regex patterns for common page markers ---
    page_number_patterns = [
        re.compile(r"^\s*page\s+\d+\s*(of\s+\d+)?\s*$", re.IGNORECASE),
        re.compile(r"^\s*-\s*\d+\s*-\s*$"),              # - 7 -
        re.compile(r"^\s*\d+\s*$"),                       # standalone number
        re.compile(r"^\s*©.*\d{4}.*$", re.IGNORECASE),    # copyright lines
    ]

    def _is_page_marker(line: str) -> bool:
        return any(p.match(line) for p in page_number_patterns)

    # --- Step 4: strip boilerplate from every page ---
    cleaned: list[str] = []
    for page_text in page_texts:
        filtered_lines = []
        for line in page_text.splitlines():
            normalized = line.strip().lower()
            if normalized in boilerplate:
                continue
            if _is_page_marker(line):
                continue
            filtered_lines.append(line)

        cleaned.append("\n".join(filtered_lines))

    return cleaned


def _ocr_pdf_page(filepath: Path, page_index: int) -> str:
    """OCR a single PDF page as a fallback when text extraction fails."""
    try:
        from pdf2image import convert_from_path
        import pytesseract

        # Convert just the one page to an image (1-indexed for pdf2image)
        images = convert_from_path(
            str(filepath),
            first_page=page_index + 1,
            last_page=page_index + 1,
            dpi=300,
        )

        if images:
            return pytesseract.image_to_string(images[0])
    except ImportError:
        # OCR libraries not installed — degrade gracefully
        pass
    except Exception:
        # OCR failed for some reason — don't crash the pipeline
        pass

    return ""


def extract_pdf(filepath: Path) -> list[Document]:
    """Extract from .pdf files. One Document per page.
    
    Strips repeated headers/footers and page-number boilerplate
    before returning, so downstream chunking gets clean text.
    """
    reader = PdfReader(str(filepath))

    # First pass: extract raw text from all pages
    raw_texts: list[str] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if len(text.strip()) < 50:
            ocr_text = _ocr_pdf_page(filepath, i)
            if ocr_text.strip():
                text = ocr_text
        raw_texts.append(text)

    # Strip boilerplate across all pages
    cleaned_texts = _strip_pdf_boilerplate(raw_texts)

    # Build Documents from cleaned text
    documents = []
    for i, text in enumerate(cleaned_texts):
        if text.strip():
            documents.append(Document(
                text=text,
                metadata={
                    "source": str(filepath),
                    "filetype": "pdf",
                    "page": i + 1,
                    "total_pages": len(reader.pages),
                },
            ))

    return documents


def extract_docx(filepath: Path) -> list[Document]:
    """Extract from .docx files, including tables in reading order."""
    doc = DocxDocument(str(filepath))

    # Walk body elements in document order to interleave
    # paragraphs and tables correctly
    PARAGRAPH_TAG = (
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
    )
    TABLE_TAG = (
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl"
    )

    # Build lookup maps: XML element -> python-docx object
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    parts: list[str] = []
    table_count = 0

    for child in doc.element.body:
        if child.tag == PARAGRAPH_TAG and child in para_map:
            text = para_map[child].text.strip()
            if text:
                parts.append(text)

        elif child.tag == TABLE_TAG and child in table_map:
            table = table_map[child]
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))
                table_count += 1

    return [Document(
        text="\n\n".join(parts),
        metadata={
            "source": str(filepath),
            "filetype": "docx",
            "paragraph_count": len(parts) - table_count,
            "table_count": table_count,
        },
    )]